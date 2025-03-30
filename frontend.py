import streamlit as st
import requests
from docx import Document
from docx.shared import Pt
from io import BytesIO

def format_resume(resume_json):
    formatted_text = f"""
**{resume_json['name']}**  
{resume_json['contact_info']}  

### **Professional Experience**
"""
    
    for exp in resume_json['experience']:
        formatted_text += f"\n**{exp['company']}** - **{exp['job_title']}** ({exp['dates']})\n"
        for achievement in exp['achievements']:
            formatted_text += f"- {achievement}\n"
    
    formatted_text += f"""
\n### **Education**\n
**{resume_json['education']['degree']}** - {resume_json['education']['year']}  

### **Skills**
"""
    for skill in resume_json['skills']:
        formatted_text += f"- {skill}\n"
    
    return formatted_text

# Function to generate properly formatted DOCX
def generate_docx(resume_json):
    doc = Document()
    doc.add_paragraph(resume_json['name'], 'Title')
    doc.add_paragraph(resume_json['contact_info'])
    
    doc.add_paragraph("Professional Experience", 'Heading1')
    for exp in resume_json['experience']:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(f"{exp['company']} - {exp['job_title']} ({exp['dates']})")
        run.bold = True
        run.font.size = Pt(12)
        
        for achievement in exp['achievements']:
            doc.add_paragraph(achievement, style="List Bullet")
    
    doc.add_paragraph("Education", 'Heading1')
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(f"{resume_json['education']['degree']} - {resume_json['education']['year']}")
    run.bold = True
    
    doc.add_paragraph("Skills", 'Heading1')
    for skill in resume_json['skills']:
        doc.add_paragraph(skill, style="List Bullet")
    
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Streamlit UI
st.title("Intelli-CV")

uploaded_file = st.file_uploader("Upload your resume (.docx)", type=["docx"])
job_description = st.text_area("Paste the job description here")

if st.button("Optimize Resume"):
    if uploaded_file and job_description:
        files = {"resume": uploaded_file}
        data = {"job_description": job_description}

        response = requests.post("https://ai-resume-writer-backend.onrender.com/optimize_resume", files=files, data=data)

        if response.status_code == 200:
            response_json = response.json()
            
            st.write("Raw AI Response:", response_json)

            import json
            optimized_resume = response_json.get("optimized_resume", "{}")  # Ensure a default empty JSON string

            # Ensure we are dealing with a dictionary
            if isinstance(optimized_resume, str):
                try:
                    optimized_resume = json.loads(optimized_resume)  # Convert to dictionary if it's a string
                except json.JSONDecodeError:
                    st.error("Error: Failed to parse AI response. Please try again.")
                    st.stop()

            # Validate that all necessary keys exist to avoid KeyErrors
            required_keys = ["name", "contact_info", "experience", "education", "skills"]
            for key in required_keys:
                if key not in optimized_resume:
                    st.error(f"Error: Missing key '{key}' in AI response.")
                    st.stop()

            
            st.subheader("Optimized Resume")
            st.markdown(format_resume(optimized_resume), unsafe_allow_html=True)

            docx_file = generate_docx(optimized_resume)
            st.download_button(
                label="Download Resume as DOCX",
                data=docx_file,
                file_name="Optimized_Resume.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
